import json as _json
import logging
import os
import tempfile

from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile
from pydantic import BaseModel

from app.middleware.rate_limiter import get_user_persona_limit, limiter
from app.routers.onboarding import get_current_user_id

logger = logging.getLogger(__name__)

router = APIRouter()

MAX_CHARS = 50000
SMART_PREVIEW_CHARS = 500
DEFAULT_MAX_UPLOAD_SIZE_BYTES = 10 * 1024 * 1024
UPLOAD_CHUNK_SIZE_BYTES = 1024 * 1024


class FileUploadResponse(BaseModel):
    filename: str
    content: str
    summary_prompt: str


class SmartUploadResponse(BaseModel):
    """Response from the smart upload endpoint with content detection."""

    filename: str
    content_type: str
    detected_type: str
    summary: str
    size_bytes: int
    suggested_actions: list[str]


def _extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text content from PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf not installed")
        return "[PDF parsing not available. Please install pypdf package.]"

    try:
        reader = PdfReader(pdf_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text)
                else:
                    logger.warning(
                        "No text found on page %s in %s", page_num + 1, pdf_path
                    )
            except Exception as exc:
                logger.warning(
                    "Error extracting text from page %s: %s", page_num + 1, exc
                )
                continue

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            return "[PDF contains no extractable text. It may contain only images or scanned documents. Please upload a text-based version or add it to Knowledge Vault manually.]"

        return full_text

    except Exception as exc:
        logger.error("Error reading PDF file %s: %s", pdf_path, exc)
        return f"[Error reading PDF: {exc!s}]"


def _extract_text_from_docx(docx_path: str) -> str:
    """Extract text content from DOCX file."""
    try:
        from docx import Document
    except ImportError:
        return "[DOCX parsing not available. Please install python-docx package.]"

    try:
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as exc:
        logger.error("Error reading DOCX file %s: %s", docx_path, exc)
        return f"[Error reading DOCX: {exc!s}]"


def _extract_text_from_text_file(file_path: str) -> str:
    """Extract text content from a plain text file."""
    encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

    for encoding in encodings:
        try:
            with open(file_path, encoding=encoding) as file_handle:
                return file_handle.read()
        except UnicodeDecodeError:
            continue
        except Exception as exc:
            logger.error("Error reading text file %s: %s", file_path, exc)
            break

    return "[Error: Could not decode file. Try saving as UTF-8.]"


def _truncate_content(content: str, max_chars: int = MAX_CHARS) -> str:
    """Truncate content if it exceeds maximum length."""
    if len(content) > max_chars:
        return (
            content[:max_chars]
            + f"\n\n... [Truncated - content exceeds {max_chars} character limit. First {max_chars} characters shown.]"
        )
    return content


def _get_max_upload_size_bytes() -> int:
    raw_value = os.getenv(
        "MAX_UPLOAD_SIZE_BYTES", str(DEFAULT_MAX_UPLOAD_SIZE_BYTES)
    ).strip()
    try:
        parsed = int(raw_value)
    except ValueError:
        logger.warning(
            "Invalid MAX_UPLOAD_SIZE_BYTES value %r; falling back to %s bytes",
            raw_value,
            DEFAULT_MAX_UPLOAD_SIZE_BYTES,
        )
        return DEFAULT_MAX_UPLOAD_SIZE_BYTES
    return parsed if parsed > 0 else DEFAULT_MAX_UPLOAD_SIZE_BYTES


def _format_bytes(num_bytes: int) -> str:
    if num_bytes >= 1024 * 1024:
        return f"{num_bytes / (1024 * 1024):.1f} MiB"
    if num_bytes >= 1024:
        return f"{num_bytes / 1024:.1f} KiB"
    return f"{num_bytes} B"


def _raise_file_too_large(max_upload_size_bytes: int) -> None:
    raise HTTPException(
        status_code=413,
        detail=f"File too large. Maximum upload size is {_format_bytes(max_upload_size_bytes)}.",
    )


def _validate_declared_upload_size(
    content_length: str | None, max_upload_size_bytes: int
) -> None:
    if not content_length:
        return

    try:
        declared_size = int(content_length)
    except ValueError:
        logger.warning(
            "Ignoring invalid Content-Length header on upload request: %r",
            content_length,
        )
        return

    if declared_size > max_upload_size_bytes:
        _raise_file_too_large(max_upload_size_bytes)


async def _write_upload_to_temp_file(
    upload: UploadFile,
    temp_path: str,
    max_upload_size_bytes: int,
) -> None:
    bytes_written = 0

    with open(temp_path, "wb") as temp_file:
        while True:
            chunk = await upload.read(UPLOAD_CHUNK_SIZE_BYTES)
            if not chunk:
                break

            bytes_written += len(chunk)
            if bytes_written > max_upload_size_bytes:
                _raise_file_too_large(max_upload_size_bytes)

            temp_file.write(chunk)


@router.post("/upload", response_model=FileUploadResponse)
@limiter.limit(get_user_persona_limit)
async def upload_file(request: Request, file: UploadFile = File(...), user_id: str = Depends(get_current_user_id)):
    """
    Upload a file, extract its text content, and return a prompt for the agent.
    """
    temp_path: str | None = None
    original_filename = file.filename or "upload"
    max_upload_size_bytes = _get_max_upload_size_bytes()

    try:
        _validate_declared_upload_size(
            request.headers.get("content-length"), max_upload_size_bytes
        )

        temp_path = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=f"_{original_filename}",
        ).name

        await _write_upload_to_temp_file(file, temp_path, max_upload_size_bytes)

        content = ""
        filename = original_filename.lower()

        if filename.endswith(
            (
                ".txt",
                ".md",
                ".csv",
                ".json",
                ".py",
                ".js",
                ".ts",
                ".html",
                ".css",
                ".sql",
                ".xml",
                ".yaml",
                ".yml",
            )
        ):
            content = _extract_text_from_text_file(temp_path)
        elif filename.endswith(".pdf"):
            content = _extract_text_from_pdf(temp_path)
        elif filename.endswith(".docx"):
            content = _extract_text_from_docx(temp_path)
        elif filename.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
            content = f"[Image file detected: {original_filename}. Image content extraction requires OCR. Please provide a text description or upload a text-based document.]"
        elif filename.endswith((".mp4", ".avi", ".mov", ".mkv", ".webm")):
            content = f"[Video file detected: {original_filename}. Video content extraction is not supported. Please provide a transcript or description.]"
        elif filename.endswith((".mp3", ".wav", ".ogg", ".flac", ".m4a")):
            content = f"[Audio file detected: {original_filename}. Audio transcription is not supported in this endpoint. Please use a dedicated transcription service.]"
        else:
            content = f"[Unsupported file type: {original_filename}. Supported types: txt, md, csv, json, py, js, ts, html, css, sql, xml, yaml, yml, pdf, docx]"

        content = _truncate_content(content)

        summary_prompt = (
            f"I have uploaded a file named '{original_filename}'. "
            f"Here is its content:\n\n"
            f"```\n{content}\n```\n\n"
            "Please analyze this file. Summarize its key points and tell me if I should add it to the Knowledge Vault."
        )

        return FileUploadResponse(
            filename=original_filename,
            content=content,
            summary_prompt=summary_prompt,
        )

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("File processing failed: %s", exc, exc_info=True)
        raise HTTPException(status_code=500, detail=f"File processing failed: {exc!s}")
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception as exc:
                logger.warning("Failed to clean up temp file %s: %s", temp_path, exc)


# ---------------------------------------------------------------------------
# Smart Upload — content-type detection + summary for Context Sniffer
# ---------------------------------------------------------------------------

# Map file extensions to (detected_type, human_label)
_EXTENSION_TYPE_MAP: dict[str, tuple[str, str]] = {
    ".pdf": ("document", "PDF Document"),
    ".csv": ("spreadsheet", "CSV Spreadsheet"),
    ".xlsx": ("spreadsheet", "Excel Spreadsheet"),
    ".xls": ("spreadsheet", "Excel Spreadsheet"),
    ".png": ("image", "PNG Image"),
    ".jpg": ("image", "JPEG Image"),
    ".jpeg": ("image", "JPEG Image"),
    ".gif": ("image", "GIF Image"),
    ".webp": ("image", "WebP Image"),
    ".txt": ("document", "Text File"),
    ".md": ("document", "Markdown Document"),
    ".doc": ("document", "Word Document"),
    ".docx": ("document", "Word Document"),
    ".json": ("data", "JSON Data"),
    ".py": ("document", "Python Source"),
    ".js": ("document", "JavaScript Source"),
    ".ts": ("document", "TypeScript Source"),
    ".html": ("document", "HTML Document"),
    ".xml": ("document", "XML Document"),
    ".yaml": ("document", "YAML Document"),
    ".yml": ("document", "YAML Document"),
    ".sql": ("document", "SQL Script"),
}


def _detect_file_type(filename: str, content_type: str) -> tuple[str, str]:
    """Return (detected_type, human_label) for a file."""
    ext = os.path.splitext(filename.lower())[1]
    if ext in _EXTENSION_TYPE_MAP:
        return _EXTENSION_TYPE_MAP[ext]
    # Fallback: derive from MIME type
    if content_type.startswith("image/"):
        return ("image", "Image")
    if content_type.startswith("text/"):
        return ("document", "Text File")
    return ("file", "File")


def _generate_pdf_preview(temp_path: str) -> str:
    """Extract a short text preview from a PDF."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(temp_path)
        page_count = len(reader.pages)
        text_parts: list[str] = []
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
                text_parts.append(text)
                if len("".join(text_parts)) >= SMART_PREVIEW_CHARS:
                    break
            except Exception:
                continue
        preview = "".join(text_parts)[:SMART_PREVIEW_CHARS].strip()
        if preview:
            return f"{page_count}-page PDF. Preview: {preview}..."
        return f"{page_count}-page PDF (no extractable text — may be scanned)."
    except ImportError:
        return "PDF document (text preview unavailable — pypdf not installed)."
    except Exception as exc:
        logger.warning("PDF preview failed: %s", exc)
        return "PDF document."


def _generate_csv_preview(temp_path: str) -> str:
    """Read headers and first few rows of a CSV file."""
    import csv

    try:
        with open(temp_path, encoding="utf-8", errors="replace") as fh:
            reader = csv.reader(fh)
            rows: list[list[str]] = []
            for i, row in enumerate(reader):
                rows.append(row)
                if i >= 3:  # header + 3 data rows
                    break
        if not rows:
            return "Empty CSV file."
        headers = rows[0]
        sample = rows[1:] if len(rows) > 1 else []
        parts = [f"Columns ({len(headers)}): {', '.join(headers[:10])}"]
        if len(headers) > 10:
            parts[0] += f" ... +{len(headers) - 10} more"
        for idx, row in enumerate(sample):
            parts.append(f"Row {idx + 1}: {', '.join(row[:10])}")
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("CSV preview failed: %s", exc)
        return "CSV spreadsheet."


def _generate_xlsx_preview(temp_path: str) -> str:
    """Read headers and first few rows of an Excel file."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(temp_path, read_only=True, data_only=True)
        ws = wb.active
        if ws is None:
            return "Excel file with no active sheet."
        rows: list[list[str]] = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            rows.append([str(cell) if cell is not None else "" for cell in row])
            if i >= 3:
                break
        wb.close()
        if not rows:
            return "Empty Excel file."
        headers = rows[0]
        sample = rows[1:] if len(rows) > 1 else []
        parts = [f"Columns ({len(headers)}): {', '.join(headers[:10])}"]
        if len(headers) > 10:
            parts[0] += f" ... +{len(headers) - 10} more"
        for idx, row in enumerate(sample):
            parts.append(f"Row {idx + 1}: {', '.join(row[:10])}")
        sheet_count = len(wb.sheetnames) if hasattr(wb, "sheetnames") else 1
        if sheet_count > 1:
            parts.insert(0, f"{sheet_count} sheets.")
        return "\n".join(parts)
    except ImportError:
        return "Excel spreadsheet (preview unavailable — openpyxl not installed)."
    except Exception as exc:
        logger.warning("XLSX preview failed: %s", exc)
        return "Excel spreadsheet."


def _generate_image_preview(temp_path: str, filename: str) -> str:
    """Return basic image metadata."""
    import struct

    size_bytes = os.path.getsize(temp_path)
    try:
        # Try to read dimensions using struct for PNG/JPEG without extra deps
        with open(temp_path, "rb") as fh:
            header = fh.read(32)
        # PNG magic: 8 bytes, then IHDR chunk with width/height as 4-byte ints
        if header[:8] == b"\x89PNG\r\n\x1a\n":
            width, height = struct.unpack(">II", header[16:24])
            return f"Image: {filename} ({width}x{height}, {_format_bytes(size_bytes)})"
        # JPEG: look for SOF0/SOF2 markers for dimensions
        if header[:2] == b"\xff\xd8":
            with open(temp_path, "rb") as fh:
                fh.read(2)
                while True:
                    marker = fh.read(2)
                    if len(marker) < 2:
                        break
                    if marker[0] != 0xFF:
                        break
                    if marker[1] in (0xC0, 0xC2):
                        fh.read(3)  # skip length + precision
                        height, width = struct.unpack(">HH", fh.read(4))
                        return f"Image: {filename} ({width}x{height}, {_format_bytes(size_bytes)})"
                    # Skip this segment
                    seg_len = struct.unpack(">H", fh.read(2))[0]
                    fh.read(seg_len - 2)
    except Exception:
        pass
    return f"Image: {filename} ({_format_bytes(size_bytes)})"


def _generate_json_preview(temp_path: str) -> str:
    """Parse JSON and show its structure."""
    try:
        with open(temp_path, encoding="utf-8") as fh:
            data = _json.load(fh)
        if isinstance(data, dict):
            keys = list(data.keys())[:10]
            extra = f" ... +{len(data) - 10} more keys" if len(data) > 10 else ""
            return f"JSON object with {len(data)} keys: {', '.join(keys)}{extra}"
        if isinstance(data, list):
            sample_type = type(data[0]).__name__ if data else "empty"
            return f"JSON array with {len(data)} items (first item type: {sample_type})"
        return f"JSON value: {str(data)[:200]}"
    except Exception as exc:
        logger.warning("JSON preview failed: %s", exc)
        return "JSON data file."


def _generate_text_preview(temp_path: str) -> str:
    """Read first N characters of a text-based file."""
    try:
        with open(temp_path, encoding="utf-8", errors="replace") as fh:
            preview = fh.read(SMART_PREVIEW_CHARS).strip()
        if not preview:
            return "Empty text file."
        suffix = "..." if len(preview) >= SMART_PREVIEW_CHARS else ""
        return f"{preview}{suffix}"
    except Exception as exc:
        logger.warning("Text preview failed: %s", exc)
        return "Text file."


def _generate_docx_preview(temp_path: str) -> str:
    """Extract a short text preview from a DOCX file."""
    try:
        from docx import Document

        doc = Document(temp_path)
        parts: list[str] = []
        total = 0
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
                total += len(parts[-1])
                if total >= SMART_PREVIEW_CHARS:
                    break
        if not parts:
            return "Word document (no text content)."
        preview = " ".join(parts)[:SMART_PREVIEW_CHARS]
        return f"{preview}..."
    except ImportError:
        return "Word document (preview unavailable — python-docx not installed)."
    except Exception as exc:
        logger.warning("DOCX preview failed: %s", exc)
        return "Word document."


def _build_smart_summary(
    temp_path: str,
    filename: str,
    detected_type: str,
    human_label: str,
) -> str:
    """Dispatch to the right preview generator based on detected type."""
    ext = os.path.splitext(filename.lower())[1]

    if ext == ".pdf":
        return _generate_pdf_preview(temp_path)
    if ext == ".csv":
        return _generate_csv_preview(temp_path)
    if ext in (".xlsx", ".xls"):
        return _generate_xlsx_preview(temp_path)
    if detected_type == "image":
        return _generate_image_preview(temp_path, filename)
    if ext == ".json":
        return _generate_json_preview(temp_path)
    if ext in (".docx", ".doc"):
        return _generate_docx_preview(temp_path)
    if detected_type == "document":
        return _generate_text_preview(temp_path)

    # Generic fallback
    return f"{human_label}: {filename} ({_format_bytes(os.path.getsize(temp_path))})"


@router.post("/upload/smart", response_model=SmartUploadResponse)
@limiter.limit(get_user_persona_limit)
async def smart_upload(request: Request, file: UploadFile = File(...), user_id: str = Depends(get_current_user_id)):
    """Smart file upload with content-type detection and preview summary.

    Returns metadata about the uploaded file so the frontend can offer
    the user a choice: add it to the Knowledge Vault, analyze it now,
    or summarize it.
    """
    temp_path: str | None = None
    original_filename = file.filename or "upload"
    mime_type = file.content_type or "application/octet-stream"
    max_upload_size_bytes = _get_max_upload_size_bytes()

    try:
        _validate_declared_upload_size(
            request.headers.get("content-length"),
            max_upload_size_bytes,
        )

        temp_path = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=f"_{original_filename}",
        ).name

        await _write_upload_to_temp_file(file, temp_path, max_upload_size_bytes)

        size_bytes = os.path.getsize(temp_path)
        detected_type, human_label = _detect_file_type(original_filename, mime_type)
        summary = _build_smart_summary(
            temp_path,
            original_filename,
            detected_type,
            human_label,
        )

        # All file types support these actions
        suggested_actions = ["add_to_vault", "analyze_now", "summarize"]

        return SmartUploadResponse(
            filename=original_filename,
            content_type=mime_type,
            detected_type=detected_type,
            summary=summary,
            size_bytes=size_bytes,
            suggested_actions=suggested_actions,
        )

    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Smart upload processing failed: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Smart upload processing failed: {exc!s}",
        )
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception as exc:
                logger.warning(
                    "Failed to clean up temp file %s: %s",
                    temp_path,
                    exc,
                )
